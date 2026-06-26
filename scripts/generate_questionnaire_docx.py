#!/usr/bin/env python3
"""Generate questionnaire docx from personality-quantitative-tester.html data."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "人格-量化因子适配测试问卷.docx"

OPTION_TEXTS = [
    ("A", "非常不同意", 1),
    ("B", "不同意", 2),
    ("C", "中立", 3),
    ("D", "同意", 4),
    ("E", "非常同意", 5),
]

DIMENSION_NAMES = {
    "openness": "开放性",
    "conscientiousness": "尽责性",
    "extraversion": "外倾性",
    "agreeableness": "宜人性",
    "neuroticism": "神经质",
    "riskTolerance": "风险承受力",
    "lossAversion": "损失厌恶度",
    "patience": "持有耐心",
    "overconfidence": "过度自信",
    "lie": "测谎题（不计入人格得分）",
}

QUESTIONS = [
    {"id": 1, "content": "我愿意尝试新的投资思路和小众策略，不局限于传统方法", "dimension": "openness", "isReverse": False},
    {"id": 2, "content": "做投资决策前，我会提前制定好明确的买入、卖出计划", "dimension": "conscientiousness", "isReverse": False},
    {"id": 3, "content": "我经常和朋友、社群交流投资信息，喜欢跟着热点方向布局", "dimension": "extraversion", "isReverse": False},
    {"id": 4, "content": "持仓出现浮亏时，我很容易焦虑不安，甚至影响睡眠", "dimension": "neuroticism", "isReverse": False},
    {"id": 5, "content": "当身边人都在买某只股票时，我很难坚持自己原本的计划", "dimension": "agreeableness", "isReverse": False},
    {"id": 6, "content": "我更喜欢成熟稳定的投资标的，不喜欢不确定性强的新事物", "dimension": "openness", "isReverse": True},
    {"id": 7, "content": "我买卖经常凭当下感觉，不会提前做详细计划", "dimension": "conscientiousness", "isReverse": True},
    {"id": 8, "content": "我更习惯自己独立研究，很少主动和别人讨论投资标的", "dimension": "extraversion", "isReverse": True},
    {"id": 9, "content": "面对亏损我能保持平静，不会因为短期波动乱了节奏", "dimension": "neuroticism", "isReverse": True},
    {"id": 10, "content": "即使和多数人意见不同，我也敢做出逆向的投资决策", "dimension": "agreeableness", "isReverse": True},
    {"id": 11, "content": "为了追求更高收益，我可以接受本金短期回撤20%以上", "dimension": "riskTolerance", "isReverse": False},
    {"id": 12, "content": "同样金额的亏损带来的痛苦，远大于同等盈利带来的快乐", "dimension": "lossAversion", "isReverse": False},
    {"id": 13, "content": "我愿意持有一只标的3年以上，来换取长期的收益回报", "dimension": "patience", "isReverse": False},
    {"id": 14, "content": "我觉得自己的选股/择时能力，超过市场上大多数投资者", "dimension": "overconfidence", "isReverse": False},
    {"id": 15, "content": "比起高收益高波动，我更愿意赚稳健、波动小的钱", "dimension": "riskTolerance", "isReverse": True},
    {"id": 16, "content": "看到持仓浮亏，我会比看到浮盈更在意、更频繁地看盘", "dimension": "lossAversion", "isReverse": False},
    {"id": 17, "content": "持有标的超过半年不涨，我就会忍不住想换仓", "dimension": "patience", "isReverse": True},
    {"id": 18, "content": "我经常因为觉得自己判断正确，而重仓押注单一标的", "dimension": "overconfidence", "isReverse": False},
    {"id": 19, "content": "我做投资以来，从来没有过任何一次后悔的操作", "dimension": "lie", "isReverse": False},
    {"id": 20, "content": "面对亏损，我从来不会有任何情绪波动", "dimension": "lie", "isReverse": False},
]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("人格-量化因子适配测试问卷", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "本问卷共 20 题，基于大五人格模型与行为金融学设计，用于评估投资者人格特质，"
        "并匹配 8 大量化因子适配度。每题请根据实际情况选择最符合的一项。",
    )
    doc.add_paragraph()

    add_heading(doc, "一、作答说明", 1)
    add_paragraph(doc, "1. 每题均为单选题，从以下 5 个选项中选择一项：")
    for letter, text, score in OPTION_TEXTS:
        add_paragraph(doc, f"   {letter}. {text}（计 {score} 分）")
    add_paragraph(doc, "2. 部分题目为反向题，系统会自动反转计分（选「非常同意」会按 1 分计）。")
    add_paragraph(doc, "3. 第 19、20 题为测谎题，若两题均选「同意」或「非常同意」，系统将判定为无效作答。")
    add_paragraph(doc, "4. 本测试仅供参考，不构成任何投资建议。")
    doc.add_paragraph()

    add_heading(doc, "二、人格维度说明", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "维度"
    headers[1].text = "说明"
    headers[2].text = "对应题号"

    dimension_questions = {}
    for q in QUESTIONS:
        if q["dimension"] == "lie":
            continue
        dimension_questions.setdefault(q["dimension"], []).append(str(q["id"]))

    dimension_desc = {
        "openness": "对新策略、新事物的接受程度",
        "conscientiousness": "投资计划性与纪律性",
        "extraversion": "社交交流与跟热点倾向",
        "agreeableness": "是否容易随大流",
        "neuroticism": "面对亏损时的情绪稳定性",
        "riskTolerance": "对波动和回撤的承受能力",
        "lossAversion": "对亏损的心理敏感度",
        "patience": "长期持有的耐心程度",
        "overconfidence": "对自身投资能力的自信程度",
    }

    for key, name in DIMENSION_NAMES.items():
        if key == "lie":
            continue
        row = table.add_row().cells
        row[0].text = name
        row[1].text = dimension_desc[key]
        row[2].text = "、".join(dimension_questions.get(key, []))

    doc.add_paragraph()

    add_heading(doc, "三、问卷题目与选项", 1)

    for q in QUESTIONS:
        add_heading(doc, f"第 {q['id']} 题", 2)

        p = doc.add_paragraph()
        run = p.add_run(q["content"])
        run.bold = True
        run.font.size = Pt(12)

        meta = doc.add_paragraph()
        dim_name = DIMENSION_NAMES[q["dimension"]]
        reverse_text = "是（计分会反转：得分 = 6 - 原始分）" if q["isReverse"] else "否"
        meta.add_run(f"所属维度：{dim_name}    |    反向题：{reverse_text}").font.size = Pt(10)
        meta.runs[0].font.color.rgb = RGBColor(0x4E, 0x59, 0x69)

        doc.add_paragraph("选项：")
        for letter, text, score in OPTION_TEXTS:
            add_paragraph(doc, f"  {letter}. {text}")

        doc.add_paragraph()

    add_heading(doc, "四、测谎题说明", 1)
    add_paragraph(doc, "第 19 题：我做投资以来，从来没有过任何一次后悔的操作")
    add_paragraph(doc, "第 20 题：面对亏损，我从来不会有任何情绪波动")
    add_paragraph(
        doc,
        "若以上两题均选择「同意（4 分）」或「非常同意（5 分）」，系统会提示「无效作答」，"
        "要求用户重新填写。",
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
